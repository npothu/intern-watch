"""Renderer output: structure, fonts, tab stops, borders — checked by
re-opening the .docx and inspecting the XML python-docx round-trips."""

from docx import Document
from docx.oxml.ns import qn

from src.resume import jd, render, select
from src.resume.bank import load_bank
from src.resume.spec import CONTENT_W_TW, FONT, MARGIN_TW, PAGE_H_TW, PAGE_W_TW

BANK = load_bank("tests/fixtures/resume_bank.json")


def _build(fixtures, tmp_path, name="jd_backend_intern.txt"):
    profile = jd.analyze((fixtures / name).read_text(encoding="utf-8"))
    plan = select.build_plan(BANK, profile)
    out = tmp_path / "out.docx"
    render.render(plan, out)
    return plan, Document(str(out))


def test_opens_and_has_all_sections(fixtures, tmp_path):
    plan, doc = _build(fixtures, tmp_path)
    text = "\n".join(p.text for p in doc.paragraphs)
    assert plan.header_name in text
    for header in ("Education", "Programming Projects", "Community"):
        assert header in text
    for proj in plan.projects:
        assert proj.name + " | " in text


def test_page_geometry(fixtures, tmp_path):
    _, doc = _build(fixtures, tmp_path)
    s = doc.sections[0]
    assert s.page_width.twips == PAGE_W_TW
    assert s.page_height.twips == PAGE_H_TW
    assert s.left_margin.twips == s.right_margin.twips == MARGIN_TW
    assert s.top_margin.twips == s.bottom_margin.twips == MARGIN_TW


def test_font_is_times_throughout(fixtures, tmp_path):
    _, doc = _build(fixtures, tmp_path)
    assert doc.styles["Normal"].font.name == FONT
    for p in doc.paragraphs:
        for r in p.runs:
            assert r.font.name in (None, FONT)


def test_dated_lines_use_single_right_tab(fixtures, tmp_path):
    plan, doc = _build(fixtures, tmp_path)
    heading = next(p for p in doc.paragraphs
                   if p.text.startswith(plan.projects[0].name))
    stops = heading.paragraph_format.tab_stops
    assert len(stops) == 1
    stop = stops[0]
    assert stop.position.twips == CONTENT_W_TW
    assert stop.alignment.name == "RIGHT"      # WD_TAB_ALIGNMENT.RIGHT
    # exactly one tab in the text, right before the date
    assert heading.text.count("\t") == 1
    assert heading.text.endswith("\t" + plan.projects[0].date)


def test_section_headers_have_bottom_border_not_underlined_tabs(fixtures,
                                                                tmp_path):
    _, doc = _build(fixtures, tmp_path)
    headers = [p for p in doc.paragraphs
               if p.text in ("Education", "Programming Projects", "Community")]
    assert len(headers) == 3
    for p in headers:
        assert "\t" not in p.text          # the old underlined-tab hack
        pbdr = p._p.pPr.find(qn("w:pBdr"))
        assert pbdr is not None
        assert pbdr.find(qn("w:bottom")) is not None


def test_blank_line_before_community_section(fixtures, tmp_path):
    _, doc = _build(fixtures, tmp_path)
    paras = doc.paragraphs
    idx = next(i for i, p in enumerate(paras) if p.text == "Community")
    assert paras[idx - 1].text == ""            # separator paragraph
    assert paras[idx - 2].text.startswith("●")  # last project bullet


def test_bullets_have_hanging_indent(fixtures, tmp_path):
    _, doc = _build(fixtures, tmp_path)
    bullets = [p for p in doc.paragraphs if p.text.startswith("●")]
    assert bullets
    for p in bullets:
        pf = p.paragraph_format
        assert pf.left_indent.twips == 720
        assert pf.first_line_indent.twips == -360
        assert p.text.startswith("●\t")


def test_hyperlinks_present(fixtures, tmp_path):
    plan, doc = _build(fixtures, tmp_path)
    rels = doc.part.rels
    urls = {r.target_ref for r in rels.values()
            if r.reltype.endswith("/hyperlink")}
    assert {link["url"] for link in plan.links} <= urls


def test_no_spacing_surprises(fixtures, tmp_path):
    """Every paragraph pins its spacing; nothing inherits Word's defaults."""
    _, doc = _build(fixtures, tmp_path)
    for p in doc.paragraphs:
        pf = p.paragraph_format
        assert pf.space_before.pt == 0
        assert pf.line_spacing == 1.0

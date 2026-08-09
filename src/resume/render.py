"""Render a ResumePlan to .docx with python-docx.

Two deliberate departures from the old node builder, both fixing its
formatting bugs:

* Dates are right-aligned with a SINGLE right tab stop at the right text
  edge (one \\t in the run). The old "\\t\\t" against a ladder of left stops
  drifted whenever a heading ran long.
* Section-header rules are paragraph bottom borders, not underlined tab
  characters — they render identically in Word, Google Docs and LibreOffice.

Bullets are literal "●<tab>" with a hanging indent (ATS-safe, no
numbering.xml part needed).
"""

from __future__ import annotations

from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_TAB_ALIGNMENT
from docx.opc.constants import RELATIONSHIP_TYPE as RT
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Pt, RGBColor, Twips

from .bank import HeadingRun
from .select import ResumePlan
from .spec import (
    BULLET_CHAR,
    BULLET_HANG_TW,
    BULLET_INDENT_TW,
    FONT,
    LINK_COLOR,
    MARGIN_TW,
    PAGE_H_TW,
    PAGE_W_TW,
    RIGHT_TAB_TW,
    SECTION_BORDER_SZ8,
    SECTION_SPACE_AFTER_PT,
    SZ_BODY,
    SZ_CONTACT,
    SZ_EDU,
    SZ_NAME,
    SZ_SECTION,
)


def _para(doc, *, center=False, space_after_pt=0.0):
    p = doc.add_paragraph()
    pf = p.paragraph_format
    pf.space_before = Pt(0)
    pf.space_after = Pt(space_after_pt)
    pf.line_spacing = 1.0
    if center:
        pf.alignment = WD_ALIGN_PARAGRAPH.CENTER
    return p


def _run(p, text, sz, *, bold=False, italics=False, underline=False,
         color=None):
    r = p.add_run(text)
    r.font.name = FONT
    r.font.size = Pt(sz)
    r.bold = bold
    r.italic = italics
    r.underline = underline
    if color:
        r.font.color.rgb = RGBColor.from_string(color)
    return r


def _right_tab(p):
    p.paragraph_format.tab_stops.add_tab_stop(
        Twips(RIGHT_TAB_TW), WD_TAB_ALIGNMENT.RIGHT)


def _hyperlink(p, text, url, sz):
    r_id = p.part.relate_to(url, RT.HYPERLINK, is_external=True)
    h = OxmlElement("w:hyperlink")
    h.set(qn("r:id"), r_id)
    r = OxmlElement("w:r")
    rpr = OxmlElement("w:rPr")
    fonts = OxmlElement("w:rFonts")
    for attr in ("w:ascii", "w:hAnsi", "w:cs"):
        fonts.set(qn(attr), FONT)
    rpr.append(fonts)
    color = OxmlElement("w:color")
    color.set(qn("w:val"), LINK_COLOR)
    rpr.append(color)
    size = OxmlElement("w:sz")
    size.set(qn("w:val"), str(int(sz * 2)))
    rpr.append(size)
    u = OxmlElement("w:u")
    u.set(qn("w:val"), "single")
    rpr.append(u)
    r.append(rpr)
    t = OxmlElement("w:t")
    t.text = text
    r.append(t)
    h.append(r)
    p._p.append(h)


def _bottom_border(p):
    """Insert w:pBdr at its schema position inside w:pPr."""
    ppr = p._p.get_or_add_pPr()
    pbdr = OxmlElement("w:pBdr")
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), str(SECTION_BORDER_SZ8))
    bottom.set(qn("w:space"), "1")
    bottom.set(qn("w:color"), "auto")
    pbdr.append(bottom)
    for tag in ("w:shd", "w:tabs", "w:spacing", "w:ind", "w:jc"):
        anchor = ppr.find(qn(tag))
        if anchor is not None:
            anchor.addprevious(pbdr)
            return
    ppr.append(pbdr)


def _section_header(doc, text):
    p = _para(doc, space_after_pt=SECTION_SPACE_AFTER_PT)
    _run(p, text, SZ_SECTION, bold=True)
    _bottom_border(p)


def _dated_line(doc, runs: list[HeadingRun], date: str, sz: float):
    p = _para(doc)
    _right_tab(p)
    for r in runs:
        _run(p, r.text, sz, bold=r.bold, italics=r.italics)
    _run(p, "\t" + date, sz)
    return p


def _bullet(doc, text):
    p = _para(doc)
    pf = p.paragraph_format
    pf.left_indent = Twips(BULLET_INDENT_TW)
    pf.first_line_indent = Twips(-BULLET_HANG_TW)
    pf.tab_stops.add_tab_stop(Twips(BULLET_INDENT_TW), WD_TAB_ALIGNMENT.LEFT)
    _run(p, f"{BULLET_CHAR}\t{text}", SZ_BODY)


def _separator(doc):
    _para(doc)
    # an empty paragraph inherits Normal's size; pin it to body size
    doc.paragraphs[-1].add_run().font.size = Pt(SZ_BODY)


def _entry(doc, entry, *, last: bool):
    _dated_line(doc, entry.heading_runs, entry.date, SZ_BODY)
    for b in entry.bullets:
        _bullet(doc, b)
    if not last:
        _separator(doc)


def _work(doc, entry, *, last: bool):
    _dated_line(doc, [HeadingRun(text=entry.company, bold=True)],
                entry.location, SZ_BODY)
    _dated_line(doc, [HeadingRun(text=entry.role, italics=True)],
                entry.date, SZ_BODY)
    for b in entry.bullets:
        _bullet(doc, b)
    if not last:
        _separator(doc)


def render(plan: ResumePlan, out_path: str | Path) -> Path:
    doc = Document()

    normal = doc.styles["Normal"]
    normal.font.name = FONT
    normal.font.size = Pt(SZ_EDU)
    normal.paragraph_format.space_before = Pt(0)
    normal.paragraph_format.space_after = Pt(0)
    normal.paragraph_format.line_spacing = 1.0

    section = doc.sections[0]
    section.page_width = Twips(PAGE_W_TW)
    section.page_height = Twips(PAGE_H_TW)
    section.top_margin = Twips(MARGIN_TW)
    section.bottom_margin = Twips(MARGIN_TW)
    section.left_margin = Twips(MARGIN_TW)
    section.right_margin = Twips(MARGIN_TW)

    # header
    p = _para(doc, center=True)
    _run(p, plan.header_name, SZ_NAME, bold=True)
    p = _para(doc, center=True)
    _run(p, plan.contact_line, SZ_CONTACT)
    p = _para(doc, center=True)
    if plan.citizen_prefix:
        _run(p, plan.citizen_prefix, SZ_CONTACT)
    for i, link in enumerate(plan.links):
        if i:
            _run(p, " | ", SZ_CONTACT)
        _hyperlink(p, link["text"], link["url"], SZ_CONTACT)

    # education
    _section_header(doc, "Education")
    _dated_line(doc, [HeadingRun(text=plan.institution, bold=True)],
                plan.grad_date, SZ_EDU)
    for line in (plan.degree, plan.threads, plan.gpa):
        if line:
            p = _para(doc)
            _run(p, line, SZ_EDU, italics=True)
    if plan.graduate_degree_text:
        _dated_line(doc, [HeadingRun(text=plan.graduate_degree_text,
                                     italics=True)],
                    plan.graduate_degree_date, SZ_EDU)
    if plan.study_abroad_text:
        _dated_line(doc, [HeadingRun(text=plan.study_abroad_text,
                                     italics=True)],
                    plan.study_abroad_date, SZ_EDU)
    p = _para(doc)
    _run(p, "Coursework:", SZ_EDU, bold=True)
    _run(p, " " + plan.coursework, SZ_EDU)

    # work experience
    if plan.work_experience:
        _separator(doc)
        _section_header(doc, "Work Experience")
        for i, job in enumerate(plan.work_experience):
            _work(doc, job, last=i == len(plan.work_experience) - 1)

    # projects
    _separator(doc)
    _section_header(doc, "Programming Projects")
    for i, proj in enumerate(plan.projects):
        _entry(doc, proj, last=i == len(plan.projects) - 1)

    # community
    _separator(doc)
    _section_header(doc, "Community")
    for i, entry in enumerate(plan.community):
        _entry(doc, entry, last=i == len(plan.community) - 1)

    # skills
    _separator(doc)
    _section_header(doc, "Skills")
    for label, value in (("Languages:", plan.languages),
                         ("Systems & Tools:", plan.tools),
                         ("Certifications:", plan.certifications)):
        if not value:
            continue
        p = _para(doc)
        _run(p, label, SZ_EDU, bold=True)
        _run(p, " " + value, SZ_EDU)

    out_path = Path(out_path)
    out_path.parent.mkdir(parents=True, exist_ok=True)
    doc.save(str(out_path))
    return out_path
